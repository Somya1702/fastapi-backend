from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from fastapi.responses import FileResponse, JSONResponse
import os
import fitz  
import openai
import re  

app = FastAPI(
    title="SCN Reply Drafting API",
    description="Upload a Show Cause Notice (PDF) and a Word file (GPT Prompt) to generate a structured reply.",
    version="1.0.6",
    docs_url="/docs",
    redoc_url="/redoc"
)

@app.get("/")
def serve_frontend():
    return FileResponse("index.html")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF document (SCN)."""
    text = ""
    pdf_document = fitz.open(pdf_path)
    for page in pdf_document:
        text += page.get_text("text") + "\n"
    return text[:5000]  # Extract more text for legal documents

def extract_gstin(text):
    """Extract GSTIN from the Show Cause Notice."""
    gstin_pattern = r"\b\d{2}[A-Z]{5}\d{4}[A-Z]{1}\d{1}[Z]{1}[A-Z\d]{1}\b"
    matches = re.findall(gstin_pattern, text)
    return matches[0] if matches else "Not Found"

def extract_text_from_word(word_path):
    """Extracts the GPT prompt instructions from the uploaded Word file."""
    doc = Document(word_path)
    extracted_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    return extracted_text.strip()

def get_gpt_response(sc_notice_text, extracted_prompt):
    """Processes the SCN text using GPT, following instructions from the Word file."""
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": extracted_prompt},  # Use the uploaded Word file's instructions
            {"role": "user", "content": sc_notice_text}
        ]
    )
    return response.choices[0].message.content.strip() or "Not Found"

def create_word_file(response_text, assessee_name, assessee_address, gstin):
    """Generate a structured Word file for the reply to the Show Cause Notice."""
    doc = Document()

    # ✅ Letterhead (Centered, Bold, Italics)
    letterhead = doc.add_paragraph()
    letterhead_run = letterhead.add_run("<LETTERHEAD>")
    letterhead_run.bold = True
    letterhead_run.italic = True
    letterhead_run.font.size = Pt(14)
    letterhead_run.font.color.rgb = RGBColor(0, 0, 0)
    letterhead.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ✅ "To," (Left) and "Date:" (Right)
    to_date_para = doc.add_paragraph()
    to_run = to_date_para.add_run("To,")
    to_run.bold = True
    to_date_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    date_run = to_date_para.add_run("\t\t\t\t\tDate: " + datetime.today().strftime("%d-%m-%Y"))
    date_run.bold = True

    # ✅ Commissioner Address (Left-aligned)
    commissioner_para = doc.add_paragraph("Commissioner of GST & Central Excise,\nCity Name,\nState - Pin Code.")
    commissioner_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # ✅ Assessee Details (Centered)
    doc.add_paragraph()  # Empty line
    assessee_para = doc.add_paragraph()
    assessee_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    assessee_para.add_run(f"{assessee_name}\n{assessee_address}\nGSTIN: {gstin}").bold = True

    # ✅ Subject Line (Bold)
    doc.add_paragraph("\nSubject: Reply to Show Cause Notice", style="Heading 2")

    # ✅ Salutation
    doc.add_paragraph("\nSir,")

    # ✅ "BRIEF FACTS OF THE CASE" (Centered, Bold, Underlined)
    facts_heading = doc.add_paragraph()
    facts_run = facts_heading.add_run("BRIEF FACTS OF THE CASE")
    facts_run.bold = True
    facts_run.underline = True
    facts_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ✅ Case Facts (Numbered, Justified)
    case_facts_para = doc.add_paragraph()
    case_facts_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    case_facts_list = response_text.split("\n")
    for idx, fact in enumerate(case_facts_list, start=1):
        case_facts_para.add_run(f"{idx}. {fact.strip()}\n")

    # ✅ Save Word File
    file_path = "SCN_Reply.docx"
    if os.path.exists(file_path):
        os.remove(file_path)
    doc.save(file_path)
    return file_path

@app.post("/upload/")
async def upload_scn(
    pdf_file: UploadFile = File(...), 
    word_file: UploadFile = File(...)
):
    """Upload and process a Show Cause Notice and a Word file for GPT instructions."""
    pdf_path = "latest_uploaded.pdf"
    word_path = "uploaded_prompt.docx"

    # ✅ Save uploaded files
    with open(pdf_path, "wb") as f:
        f.write(await pdf_file.read())

    with open(word_path, "wb") as f:
        f.write(await word_file.read())

    extracted_text = extract_text_from_pdf(pdf_path)
    extracted_prompt = extract_text_from_word(word_path)

    # ✅ Extract Assessee Details
    assessee_name = get_gpt_response(extracted_text, "Extract the taxpayer's name from the SCN.")
    assessee_address = get_gpt_response(extracted_text, "Extract the taxpayer's address from the SCN.")
    gstin = extract_gstin(extracted_text)

    # ✅ Process SCN facts using GPT prompt
    gpt_response = get_gpt_response(extracted_text, extracted_prompt)
    word_path = create_word_file(gpt_response, assessee_name, assessee_address, gstin)

    return JSONResponse(content={
        "message": "Success",
        "gstin": gstin,
        "assessee_name": assessee_name,
        "assessee_address": assessee_address,
        "extracted_text": gpt_response,
        "download_url": "/download/"
    })

@app.get("/download/")
async def download_file():
    """Download the generated Word file for SCN reply."""
    file_path = "SCN_Reply.docx"
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="SCN_Reply.docx")
